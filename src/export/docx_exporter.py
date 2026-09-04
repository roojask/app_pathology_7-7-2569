import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from io import BytesIO
import datetime
from typing import Dict, Any, List

def set_cell_background(cell, fill_hex: str):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    """Set subtle light gray borders for table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def normalize_report_dict(source: Any) -> Dict[str, Any]:
    """
    Normalizes form data whether it originates from Flask request.form,
    a python dictionary, or JSON loaded from PostgreSQL FormHistory.
    """
    if source is None:
        source = {}

    def get_val(k: str, default: str = "") -> str:
        v = source.get(k)
        if v is None:
            return default
        return str(v).strip()

    def get_bool(k: str) -> bool:
        v = source.get(k)
        return v in [True, "true", "True", "1", 1, "on"]

    def get_list(k: str) -> List[str]:
        if hasattr(source, "getlist"):
            vals = source.getlist(k)
            if vals:
                return [str(x).strip() for x in vals if str(x).strip()]
        val = source.get(k)
        if isinstance(val, list):
            return [str(x).strip() for x in val if str(x).strip()]
        if val:
            return [str(val).strip()]
        return []

    def get_dims(k: str, count: int = 3) -> List[str]:
        if k in source and isinstance(source[k], list):
            return [str(x).strip() for x in source[k] if str(x).strip()]
        dims = []
        for i in range(count):
            d = source.get(f"{k}_{i}")
            if d is not None and str(d).strip():
                dims.append(str(d).strip())
        return dims

    data = {
        "s0_surgical_no": get_val("s0_surgical_no"),
        "s1_side": get_val("s1_side"),
        "s2_proc": get_val("s2_proc"),
        "s2_other_text": get_val("s2_other_text"),
        "s3_dims": get_dims("s3_dims", 3),
        "s4_check": get_bool("s4_check"),
        "s4_dims": get_dims("s4_dims", 3),
        "s5_dims": get_dims("s5_dims", 2),
        "s5_appears_normal": get_bool("s5_appears_normal"),
        "s6_check": get_bool("s6_check"),
        "s7_len": get_val("s7_len"),
        "s7_locs": get_list("s7_locs"),
        "s8_check": get_bool("s8_check"),
        "s8_dims": get_dims("s8_dims", 2),
        "s8_locs": get_list("s8_locs"),
        "s9_val": get_list("s9_val"),
        "s9_ulcer_text": get_val("s9_ulcer_text"),
        "s10_grammar": get_val("s10_grammar", "is a"),
        "s10_infiltrative": get_bool("s10_infiltrative"),
        "s10_inf_dims": get_dims("s10_inf_dims", 3),
        "s10_well": get_bool("s10_well"),
        "s10_well_dims": get_dims("s10_well_dims", 3),
        "s10_prev1": get_bool("s10_prev1"),
        "s10_prev1_dims": get_dims("s10_prev1_dims", 3),
        "s10_prev2": get_bool("s10_prev2"),
        "s10_prev2_cavity_dims": get_dims("s10_prev2_cavity_dims", 3),
        "s10_prev2_mass_dims": get_dims("s10_prev2_mass_dims", 3),
        "s10_5_nipple": get_bool("s10_5_nipple"),
        "s10_5_scar": get_bool("s10_5_scar"),
        "s10_5_central": get_bool("s10_5_central"),
        "s10_5_quadrant_check": get_bool("s10_5_quadrant_check"),
        "s10_5_quadrant_vals": get_list("s10_5_quadrant_vals"),
        "s10_5_other_check": get_bool("s10_5_other_check"),
        "s10_5_other": get_val("s10_5_other"),
        "s11_deep": get_val("s11_deep"),
        "s11_superior": get_val("s11_superior"),
        "s11_inferior": get_val("s11_inferior"),
        "s11_medial": get_val("s11_medial"),
        "s11_lateral": get_val("s11_lateral"),
        "s11_skin": get_val("s11_skin"),
        "s12_check": get_bool("s12_check"),
        "s12_val_left": get_val("s12_val_left"),
        "s12_val_right": get_val("s12_val_right"),
        "s13_unremarkable": get_bool("s13_unremarkable") or get_val("s13_type") == "unremarkable",
        "s13_type": get_val("s13_type"),
        "s13_text": get_val("s13_text"),
        "s14_check": get_bool("s14_check"),
        "s14_min": get_val("s14_min"),
        "s14_max": get_val("s14_max"),
        "footer_prosecutor": get_val("footer_prosecutor"),
        "footer_date": get_val("footer_date"),
    }

    # Handle sections mapping
    raw_sections = source.get("sections")
    sections_dict = {}
    section_keys = [
        ("sec_nipple", "= nipple"),
        ("sec_mass", "= mass"),
        ("sec_old_biopsy", "= old biopsy cavity with fibrosis"),
        ("sec_deep_margin", "= deep resected margin"),
        ("sec_nearest_margin", "= nearest resected margin"),
        ("sec_upper_inner", "= sampling upper inner quadrant"),
        ("sec_upper_outer", "= sampling upper outer quadrant"),
        ("sec_lower_inner", "= sampling lower inner quadrant"),
        ("sec_lower_outer", "= sampling lower outer quadrant"),
        ("sec_central", "= sampling central region"),
        ("sec_axillary", "= axillary lymph nodes")
    ]

    for form_k, label in section_keys:
        if isinstance(raw_sections, dict) and label in raw_sections:
            item = raw_sections[label]
            if isinstance(item, dict):
                sections_dict[label] = {
                    "code": str(item.get("code", "")).strip(),
                    "extra": str(item.get("extra", "")).strip()
                }
            else:
                sections_dict[label] = {"code": str(item).strip(), "extra": ""}
        else:
            code = get_val(form_k)
            extra = ""
            if "nearest" in label or "deep" in label:
                extra_k = form_k.replace("sec_", "sec_extra_")
                extra = get_val(extra_k)
            sections_dict[label] = {"code": code, "extra": extra}

    data["sections"] = sections_dict
    return data


def generate_docx_document(source_data: Any) -> BytesIO:
    """
    Generates a high-quality, professional Microsoft Word (.docx) document
    containing all 15 sections of the Pathology Gross Examination Report.
    """
    data = normalize_report_dict(source_data)

    doc = docx.Document()

    # Page Margins: Standard medical report margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    # Base styling
    normal_style = doc.styles["Normal"]
    font = normal_style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(17, 24, 39)  # Slate-900

    def format_dims(dims: List[str], unit: str = "cm") -> str:
        clean = [d for d in dims if d]
        if not clean:
            return f"_____ {unit}"
        return " x ".join(clean) + f" {unit}"

    # 1. Header: Clean Surgical Number
    raw_sno = data["s0_surgical_no"]
    clean_sno = raw_sno
    for pfx in ["Surgical Number:", "Surgical Number", "S-", "S -", "S ", "s-", "s "]:
        if clean_sno.startswith(pfx):
            clean_sno = clean_sno[len(pfx):].strip()
    formatted_sno = f"S-{clean_sno}" if clean_sno else "S-Unknown"

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_header.paragraph_format.space_after = Pt(2)
    r_sno_label = p_header.add_run("Surgical Number: ")
    r_sno_label.font.size = Pt(11)
    r_sno_label.font.color.rgb = RGBColor(75, 85, 99)
    r_sno_val = p_header.add_run(formatted_sno)
    r_sno_val.bold = True
    r_sno_val.font.size = Pt(12)
    r_sno_val.font.color.rgb = RGBColor(17, 24, 39)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_main_title = p_title.add_run("Breast gross 1\n")
    r_main_title.bold = True
    r_main_title.font.size = Pt(15)
    r_main_title.font.color.rgb = RGBColor(15, 23, 42)

    r_sub_title = p_title.add_run("Pathology Gross Examination Report")
    r_sub_title.bold = True
    r_sub_title.font.size = Pt(12)
    r_sub_title.font.color.rgb = RGBColor(71, 85, 105)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(14)
    r_div = p_div.add_run("-" * 55)
    r_div.font.color.rgb = RGBColor(203, 213, 225)
    r_div.font.size = Pt(9)

    # Helper for Section Headings
    def add_section_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(30, 58, 138)  # Deep Navy #1e3a8a
        return p

    # --- 1. Specimen Details & Procedure ---
    p_specimen = doc.add_paragraph()
    p_specimen.paragraph_format.space_after = Pt(6)
    p_specimen.paragraph_format.line_spacing = 1.2
    
    side_str = data["s1_side"] if data["s1_side"] else "__________"
    proc_val = data["s2_proc"]
    if proc_val == "modified":
        proc_str = "modified radical mastectomy"
    elif proc_val == "simple":
        proc_str = "simple mastectomy"
    elif proc_val == "other" or data["s2_other_text"]:
        proc_str = data["s2_other_text"] or "mastectomy"
    else:
        proc_str = "mastectomy"

    p_specimen.add_run("Received in formalin is a ")
    r_side = p_specimen.add_run(side_str)
    r_side.bold = True
    p_specimen.add_run(f" {proc_str} specimen.")
    if data["s2_other_text"] and proc_val not in ["other", ""]:
        p_specimen.add_run(f" ({data['s2_other_text']})")

    # Specimen Dimensions
    p_meas = doc.add_paragraph()
    p_meas.paragraph_format.space_after = Pt(4)
    p_meas.add_run("Measuring: ").bold = True
    p_meas.add_run(format_dims(data["s3_dims"]))
    
    # Axillary content
    if data["s4_check"] or data["s4_dims"]:
        p_meas.add_run(", with axillary content measuring: ")
        p_meas.add_run(format_dims(data["s4_dims"]))
    p_meas.add_run(".")

    # --- 2. Skin, Surgical Scar, & Nipple ---
    p_skin = doc.add_paragraph()
    p_skin.paragraph_format.space_after = Pt(6)
    p_skin.add_run("The skin ellipse: ").bold = True
    p_skin.add_run(format_dims(data["s5_dims"]))
    if data["s5_appears_normal"]:
        p_skin.add_run(", and appears normal.")
    else:
        p_skin.add_run(".")

    # Surgical Scar
    if data["s6_check"] or data["s7_len"] or data["s7_locs"]:
        p_scar = doc.add_paragraph()
        p_scar.paragraph_format.space_after = Pt(4)
        p_scar.paragraph_format.left_indent = Inches(0.2)
        p_scar.add_run("• Shows an old surgical scar: ").bold = True
        len_txt = data["s7_len"] if data["s7_len"] else "___"
        locs_txt = " / ".join(data["s7_locs"]) if data["s7_locs"] else "unspecified"
        p_scar.add_run(f"{len_txt} cm in length at ({locs_txt}) quadrant.")

    # Ulceration
    if data["s8_check"] or data["s8_dims"] or data["s8_locs"]:
        p_ulc = doc.add_paragraph()
        p_ulc.paragraph_format.space_after = Pt(4)
        p_ulc.paragraph_format.left_indent = Inches(0.2)
        p_ulc.add_run("• Shows an ulceration: ").bold = True
        u_dims = format_dims(data["s8_dims"])
        u_locs = " / ".join(data["s8_locs"]) if data["s8_locs"] else "unspecified"
        p_ulc.add_run(f"{u_dims} at ({u_locs}) quadrant.")

    # Nipple
    if data["s9_val"] or data["s9_ulcer_text"]:
        p_nip = doc.add_paragraph()
        p_nip.paragraph_format.space_after = Pt(6)
        p_nip.paragraph_format.left_indent = Inches(0.2)
        p_nip.add_run("• The nipple: ").bold = True
        n_vals = data["s9_val"]
        status_txt = ", ".join(n_vals) if n_vals else "evaluated"
        p_nip.add_run(f"is {status_txt}")
        if data["s9_ulcer_text"]:
            p_nip.add_run(f" ({data['s9_ulcer_text']})")
        p_nip.add_run(".")

    # --- 3. Tumor Characteristics & Location ---
    add_section_heading("Tumor Characteristics & Location")
    
    grammar = data["s10_grammar"] or "is a"
    mass_items_found = False

    if data["s10_infiltrative"] or data["s10_inf_dims"]:
        mass_items_found = True
        p_m1 = doc.add_paragraph()
        p_m1.paragraph_format.space_after = Pt(4)
        p_m1.paragraph_format.left_indent = Inches(0.2)
        p_m1.add_run(f"• There {grammar} ").bold = False
        p_m1.add_run("infiltrative firm yellow-white mass").bold = True
        p_m1.add_run(f", measuring: {format_dims(data['s10_inf_dims'])}.")

    if data["s10_well"] or data["s10_well_dims"]:
        mass_items_found = True
        p_m2 = doc.add_paragraph()
        p_m2.paragraph_format.space_after = Pt(4)
        p_m2.paragraph_format.left_indent = Inches(0.2)
        p_m2.add_run(f"• There {grammar} ").bold = False
        p_m2.add_run("well-defined firm white mass with slit-like appearance").bold = True
        p_m2.add_run(f", measuring: {format_dims(data['s10_well_dims'])}.")

    if data["s10_prev1"] or data["s10_prev1_dims"]:
        mass_items_found = True
        p_m3 = doc.add_paragraph()
        p_m3.paragraph_format.space_after = Pt(4)
        p_m3.paragraph_format.left_indent = Inches(0.2)
        p_m3.add_run(f"• There {grammar} ").bold = False
        p_m3.add_run("previous surgical cavity with adjacent fibrous tissue").bold = True
        p_m3.add_run(f", measuring: {format_dims(data['s10_prev1_dims'])}.")

    if data["s10_prev2"] or data["s10_prev2_cavity_dims"] or data["s10_prev2_mass_dims"]:
        mass_items_found = True
        p_m4 = doc.add_paragraph()
        p_m4.paragraph_format.space_after = Pt(4)
        p_m4.paragraph_format.left_indent = Inches(0.2)
        p_m4.add_run(f"• There {grammar} ").bold = False
        p_m4.add_run("previous surgical cavity with adjacent fibrous tissue").bold = True
        p_m4.add_run(f", measuring: {format_dims(data['s10_prev2_cavity_dims'])}, and a ")
        p_m4.add_run("firm yellow-white residual mass").bold = True
        p_m4.add_run(f", measuring: {format_dims(data['s10_prev2_mass_dims'])}.")

    if not mass_items_found:
        p_none = doc.add_paragraph()
        p_none.paragraph_format.space_after = Pt(4)
        p_none.paragraph_format.left_indent = Inches(0.2)
        p_none.add_run(f"• There {grammar} no dominant mass lesion identified on serial sectioning.")

    # Location
    loc_items = []
    if data["s10_5_nipple"]:
        loc_items.append("beneath the nipple")
    if data["s10_5_scar"]:
        loc_items.append("beneath the scar")
    if data["s10_5_central"]:
        loc_items.append("in the central portion (subareola)")
    if data["s10_5_quadrant_check"] or data["s10_5_quadrant_vals"]:
        q_str = " / ".join(data["s10_5_quadrant_vals"]) if data["s10_5_quadrant_vals"] else "unspecified"
        loc_items.append(f"in ({q_str}) quadrant")
    if data["s10_5_other"]:
        loc_items.append(data["s10_5_other"])

    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.space_after = Pt(6)
    p_loc.paragraph_format.left_indent = Inches(0.2)
    p_loc.add_run("• Tumor location: ").bold = True
    if loc_items:
        p_loc.add_run(", ".join(loc_items) + ".")
    else:
        p_loc.add_run("Not specified.")

    # --- 4. Resection Margins ---
    add_section_heading("Tumor and Resection Margins")

    m_deep = f"{data['s11_deep']} cm" if data["s11_deep"] else "_____ cm"
    m_sup = f"{data['s11_superior']} cm" if data["s11_superior"] else "_____ cm"
    m_inf = f"{data['s11_inferior']} cm" if data["s11_inferior"] else "_____ cm"
    m_med = f"{data['s11_medial']} cm" if data["s11_medial"] else "_____ cm"
    m_lat = f"{data['s11_lateral']} cm" if data["s11_lateral"] else "_____ cm"
    m_skin = f"{data['s11_skin']} cm" if data["s11_skin"] else "_____ cm"

    table_m = doc.add_table(rows=3, cols=2)
    table_m.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_m.autofit = False
    set_table_borders(table_m, color="E2E8F0", sz="4")

    margin_data = [
        (f"Deep margin: {m_deep}", f"Superior margin: {m_sup}"),
        (f"Inferior margin: {m_inf}", f"Medial margin: {m_med}"),
        (f"Lateral margin: {m_lat}", f"Skin margin: {m_skin}")
    ]

    for row_idx, (col1_txt, col2_txt) in enumerate(margin_data):
        row = table_m.rows[row_idx]
        row.height = Pt(22)
        for col_idx, txt in enumerate([col1_txt, col2_txt]):
            cell = row.cells[col_idx]
            cell.width = Inches(3.3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            parts = txt.split(":")
            r_k = p.add_run(parts[0] + ":")
            r_k.bold = True
            r_k.font.size = Pt(10)
            r_k.font.color.rgb = RGBColor(71, 85, 105)
            r_v = p.add_run(parts[1])
            r_v.font.size = Pt(10.5)
            r_v.font.color.rgb = RGBColor(15, 23, 42)

    p_sp1 = doc.add_paragraph()
    p_sp1.paragraph_format.space_after = Pt(2)

    # --- 5. Uninvolved Breast Parenchyma & Lymph Nodes ---
    add_section_heading("Uninvolved Breast Parenchyma & Lymph Nodes")

    p_parenchyma = doc.add_paragraph()
    p_parenchyma.paragraph_format.space_after = Pt(4)
    p_parenchyma.paragraph_format.left_indent = Inches(0.2)
    p_parenchyma.add_run("• Uninvolved parenchyma: ").bold = True
    v_left = data["s12_val_left"] if data["s12_val_left"] else "___"
    v_right = data["s12_val_right"] if data["s12_val_right"] else "___"
    p_parenchyma.add_run(f"fat to fibrous tissue ratio of approximately {v_left} : {v_right}.")

    p_rem = doc.add_paragraph()
    p_rem.paragraph_format.space_after = Pt(4)
    p_rem.paragraph_format.left_indent = Inches(0.2)
    p_rem.add_run("• Remaining breast tissue: ").bold = True
    if data["s13_unremarkable"]:
        p_rem.add_run("unremarkable.")
    elif data["s13_text"]:
        p_rem.add_run(f"{data['s13_text']}.")
    else:
        p_rem.add_run("unremarkable.")

    if data["s14_check"] or data["s14_min"] or data["s14_max"]:
        p_ln = doc.add_paragraph()
        p_ln.paragraph_format.space_after = Pt(6)
        p_ln.paragraph_format.left_indent = Inches(0.2)
        p_ln.add_run("• Lymph nodes: ").bold = True
        ln_min = data["s14_min"] if data["s14_min"] else "___"
        ln_max = data["s14_max"] if data["s14_max"] else "___"
        p_ln.add_run(f"multiple lymph nodes identified, ranging from {ln_min} cm to {ln_max} cm in diameter.")

    # --- 6. Representative Sections Submitted ---
    add_section_heading("Representative Sections Submitted")

    sections_schema = [
        ("= nipple", "Nipple"),
        ("= mass", "Mass"),
        ("= old biopsy cavity with fibrosis", "Old biopsy cavity with fibrosis"),
        ("= deep resected margin", "Deep resected margin"),
        ("= nearest resected margin", "Nearest resected margin"),
        ("= sampling upper inner quadrant", "Sampling upper inner quadrant"),
        ("= sampling upper outer quadrant", "Sampling upper outer quadrant"),
        ("= sampling lower inner quadrant", "Sampling lower inner quadrant"),
        ("= sampling lower outer quadrant", "Sampling lower outer quadrant"),
        ("= sampling central region", "Sampling central region"),
        ("= axillary lymph nodes", "Axillary lymph nodes"),
    ]

    sec_table = doc.add_table(rows=len(sections_schema) + 1, cols=2)
    sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sec_table.autofit = False
    set_table_borders(sec_table, color="CBD5E1", sz="4")

    # Table Header Row
    hdr_row = sec_table.rows[0]
    hdr_row.height = Pt(22)
    set_cell_background(hdr_row.cells[0], "F1F5F9")
    set_cell_background(hdr_row.cells[1], "F1F5F9")

    hdr_cell_0 = hdr_row.cells[0]
    hdr_cell_0.width = Inches(1.8)
    set_cell_margins(hdr_cell_0, top=100, bottom=100, left=140, right=140)
    p_h0 = hdr_cell_0.paragraphs[0]
    p_h0.paragraph_format.space_after = Pt(0)
    r_h0 = p_h0.add_run("Block / Section ID")
    r_h0.bold = True
    r_h0.font.size = Pt(10)
    r_h0.font.color.rgb = RGBColor(51, 65, 85)

    hdr_cell_1 = hdr_row.cells[1]
    hdr_cell_1.width = Inches(4.8)
    set_cell_margins(hdr_cell_1, top=100, bottom=100, left=140, right=140)
    p_h1 = hdr_cell_1.paragraphs[0]
    p_h1.paragraph_format.space_after = Pt(0)
    r_h1 = p_h1.add_run("Anatomical Sampling Site")
    r_h1.bold = True
    r_h1.font.size = Pt(10)
    r_h1.font.color.rgb = RGBColor(51, 65, 85)

    # Fill Sections Rows
    for idx, (anchor, label) in enumerate(sections_schema, start=1):
        row = sec_table.rows[idx]
        row.height = Pt(19)
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(4.8)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(c0, top=60, bottom=60, left=140, right=140)
        set_cell_margins(c1, top=60, bottom=60, left=140, right=140)

        if idx % 2 == 1:
            set_cell_background(c0, "FAFAFA")
            set_cell_background(c1, "FAFAFA")

        item = data["sections"].get(anchor, {})
        code_val = item.get("code", "") if isinstance(item, dict) else str(item)
        extra_val = item.get("extra", "") if isinstance(item, dict) else ""

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        if code_val:
            r0 = p0.add_run(code_val)
            r0.bold = True
            r0.font.size = Pt(10.5)
            r0.font.color.rgb = RGBColor(29, 78, 216)  # Blue-700
        else:
            r0 = p0.add_run("_______")
            r0.font.size = Pt(9.5)
            r0.font.color.rgb = RGBColor(203, 213, 225)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(f"= {label}")
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(30, 41, 59)
        if extra_val:
            r_extra = p1.add_run(f" ({extra_val})")
            r_extra.font.size = Pt(9.5)
            r_extra.font.italic = True
            r_extra.font.color.rgb = RGBColor(100, 116, 139)

    # --- 7. Footer Sign-Off ---
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(24)
    p_footer.paragraph_format.space_after = Pt(0)
    p_footer.paragraph_format.line_spacing = 1.3

    prosecutor = data["footer_prosecutor"] or "................................................."
    date_val = data["footer_date"] or datetime.datetime.now().strftime("%d/%m/%Y")

    r_f1 = p_footer.add_run(f"Prosecutor: {prosecutor}\n")
    r_f1.font.size = Pt(10.5)
    r_f1.font.color.rgb = RGBColor(51, 65, 85)

    r_f2 = p_footer.add_run(f"Date: {date_val}\n")
    r_f2.font.size = Pt(10.5)
    r_f2.font.color.rgb = RGBColor(51, 65, 85)

    r_f3 = p_footer.add_run("Approved in conference")
    r_f3.bold = True
    r_f3.font.size = Pt(10)
    r_f3.font.color.rgb = RGBColor(15, 23, 42)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
